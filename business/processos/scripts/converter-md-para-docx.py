import sys
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_inline(paragraph, text):
    # handle **bold** and *italic* and `code`
    pattern = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 1:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
        else:
            paragraph.add_run(part)

def convert(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().splitlines()

    doc = Document()
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped == '' :
            i += 1
            continue

        if stripped == '---':
            doc.add_paragraph('_' * 40)
            i += 1
            continue

        # table detection
        if stripped.startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [ [c.strip() for c in tl.strip('|').split('|')] for tl in table_lines ]
            # remove separator row (---)
            rows = [r for r in rows if not all(re.match(r'^:?-+:?$', c) for c in r)]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        if c_idx < len(table.rows[r_idx].cells):
                            table.rows[r_idx].cells[c_idx].text = cell
            continue

        # headers
        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # blockquote
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_inline(p, stripped.lstrip('> ').strip())
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # bullet list
        if stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, stripped[2:].strip())
            i += 1
            continue

        # numbered list
        m2 = re.match(r'^\d+\.\s+(.*)$', stripped)
        if m2:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, m2.group(1))
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    doc.save(docx_path)
    print(f"OK: {docx_path}")

if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
